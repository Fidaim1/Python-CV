from telnetlib import DO
from turtle import width
from docx import Document
from docx.shared import Inches

Document = Document ()

# Profile picture

Document.add_picture(
     'fidaa.jpg',
     width=Inches(2.0)
)

# Name , Phone, Email

Document.add_heading('Personal Details')

name = 'Name: Fidaim Ahmeti'
phone_number = 'Phone: +38349844721'
email = 'Email: fidaimahmeti1@gmail.com'

Document.add_paragraph(
    name + ' | ' + phone_number + ' | ' + email)

# Skills

Document.add_heading('Skills')
my_skills = ('Microsoft Office | Coding | Marketing | Social Media')

Document.add_paragraph(my_skills)

# Languages

Document.add_heading('Languages')
my_language = ('English')

Document.add_paragraph(my_language)

# Education 

Document.add_heading('Education')
p = Document.add_paragraph()

university = ('University of Prishtina "Hasan Prishtina" ')
from_date = ('Oct 2018 ')
to_date = ('Jul 2021 ')

p.add_run(university + ' ').bold = True
p.add_run(from_date + '-' + to_date + '\n').italic = True

department = ('Master Degree in Marketing')
p.add_run(department).bold = True

# Internships

Document.add_heading('Internships')
p = Document.add_paragraph()

place = ('Customs Officer ')
from_date = ('Jun 2019 ')
to_date = ('Aug 2019 ')

p.add_run(place + ' ').bold = True
p.add_run(from_date + '-' + to_date + '\n').italic = True

describe_w = ('During this internship as a customs officer I had as a duty to combat the importation of illegal goods, firearms, drugs or other dangerous of illegal items while checking the legality of items brought across national borders.')
p.add_run(describe_w)

# Courses

Document.add_heading('Courses')
p = Document.add_paragraph()

course = ('Python Programming Language ')
from_date = ('Nov 2021 ')
to_date = ('Jan 2021 ')

p.add_run(course + ' ').bold = True
p.add_run(from_date + '-' + to_date + '\n').italic = True

describe = ('During this training I have learned all about Python programming language. I like Python because has a simple syntax to code.')
p.add_run(describe)

# footer

section = Document.sections[0]
footer = section.footer
p = footer.paragraphs[0]
p.text = ('CV generated using Python Language')




Document.save('cv.docx')